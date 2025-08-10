import PyPDF2
import docx
from typing import Optional, Dict, Any
import logging
import os

logger = logging.getLogger(__name__)

class DocumentParser:
    """Service for parsing different document formats and extracting text"""
    
    def __init__(self):
        self.supported_formats = {
            'application/pdf': self._parse_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._parse_docx,
            'application/msword': self._parse_docx,
        }
    
    def extract_text(self, file_path: str, mime_type: str) -> Optional[str]:
        """
        Extract text from document based on MIME type
        
        Args:
            file_path: Path to the document file
            mime_type: MIME type of the document
            
        Returns:
            Extracted text or None if parsing failed
        """
        try:
            if not os.path.exists(file_path):
                logger.error(f"File not found: {file_path}")
                return None
            
            parser_func = self.supported_formats.get(mime_type)
            if not parser_func:
                logger.error(f"Unsupported MIME type: {mime_type}")
                return None
            
            logger.info(f"Parsing document: {file_path} (type: {mime_type})")
            return parser_func(file_path)
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF document and extract text"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"[PAGE {page_num + 1}]\n{page_text}")
                    except Exception as page_error:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {page_error}")
                        continue
                
                if not text_content:
                    logger.warning(f"No text content extracted from PDF: {file_path}")
                    return ""
                
                extracted_text = "\n\n".join(text_content)
                logger.info(f"Successfully extracted {len(extracted_text)} characters from PDF")
                return extracted_text
                
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            raise
    
    def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX document and extract text"""
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            if not text_content:
                logger.warning(f"No text content extracted from DOCX: {file_path}")
                return ""
            
            extracted_text = "\n".join(text_content)
            logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise
    
    def get_document_metadata(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """
        Extract metadata from document
        
        Returns:
            Dictionary containing document metadata
        """
        metadata = {
            'file_size': os.path.getsize(file_path),
            'mime_type': mime_type,
            'parser_version': '1.0'
        }
        
        try:
            if mime_type == 'application/pdf':
                metadata.update(self._get_pdf_metadata(file_path))
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                metadata.update(self._get_docx_metadata(file_path))
        except Exception as e:
            logger.warning(f"Error extracting metadata from {file_path}: {str(e)}")
        
        return metadata
    
    def _get_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract PDF-specific metadata"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata = {
                    'page_count': len(pdf_reader.pages),
                    'encrypted': pdf_reader.is_encrypted
                }
                
                # Extract PDF info if available
                if pdf_reader.metadata:
                    pdf_info = pdf_reader.metadata
                    metadata.update({
                        'title': pdf_info.get('/Title', ''),
                        'author': pdf_info.get('/Author', ''),
                        'creator': pdf_info.get('/Creator', ''),
                        'producer': pdf_info.get('/Producer', ''),
                        'creation_date': str(pdf_info.get('/CreationDate', '')),
                        'modification_date': str(pdf_info.get('/ModDate', ''))
                    })
                
                return metadata
        except Exception as e:
            logger.warning(f"Error extracting PDF metadata: {str(e)}")
            return {}
    
    def _get_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract DOCX-specific metadata"""
        try:
            doc = docx.Document(file_path)
            properties = doc.core_properties
            
            metadata = {
                'title': properties.title or '',
                'author': properties.author or '',
                'created': str(properties.created) if properties.created else '',
                'modified': str(properties.modified) if properties.modified else '',
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
            
            return metadata
        except Exception as e:
            logger.warning(f"Error extracting DOCX metadata: {str(e)}")
            return {}